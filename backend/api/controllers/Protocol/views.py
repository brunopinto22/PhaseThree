"""
REQ-7: Protocol Generation endpoints
Generates and manages protocol documents for approved placements.
"""
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from rest_framework.status import *
from django.http import FileResponse
from django.db import transaction
from django.utils import timezone
from django.core.files.base import ContentFile
from docx import Document
from io import BytesIO
import os

from api.models import (
    Candidature, CandidatureProposal, CandidatureHistory, Protocol,
    Student, Proposal, Teacher, Calendar, Accounts
)
from api.token_manager import decode_token


def get_user_account(auth_header):
    """Helper to decode token and get user account."""
    user_id, user_email, user_type = decode_token(auth_header)
    
    if user_email in ["Expired Token.", "Invalid Token", "Payload does not contain 'user_id'."]:
        return None, None, user_email
    
    try:
        account = Accounts.objects.get(email=user_email)
        return account, user_type, None
    except Accounts.DoesNotExist:
        return None, None, "User not found"


@api_view(["POST"])
def generateProtocol(request, pk):
    """
    REQ-7: Generate protocol document for an approved candidature.
    Only callable when candidature state is 'placed'.
    """
    auth_header = request.headers.get("Authorization")
    account, user_type, error = get_user_account(auth_header)
    
    if error:
        return Response({"message": "login"}, status=HTTP_401_UNAUTHORIZED)
    
    # Only admin, teachers (commission), or academic services can generate protocols
    if user_type not in ["admin", "teacher"]:
        return Response({"message": "Não tem permissão para gerar protocolos"}, status=HTTP_403_FORBIDDEN)
    
    try:
        candidature = Candidature.objects.get(id_candidature=pk)
        
        # Check if candidature is in the right state
        if candidature.state != 'placed':
            return Response(
                {"message": f"Candidatura deve estar no estado 'placed' para gerar protocolo. Estado atual: {candidature.state}"},
                status=HTTP_400_BAD_REQUEST
            )
        
        # Check if protocol already exists
        if hasattr(candidature, 'protocol') and candidature.protocol:
            return Response(
                {"message": "Protocolo já foi gerado para esta candidatura", "protocol_id": candidature.protocol.id_protocol},
                status=HTTP_400_BAD_REQUEST
            )
        
        # Get the accepted proposal
        accepted_cp = CandidatureProposal.objects.filter(
            candidature=candidature,
            state='accepted'
        ).first()
        
        if not accepted_cp:
            return Response(
                {"message": "Nenhuma proposta aceite encontrada para esta candidatura"},
                status=HTTP_400_BAD_REQUEST
            )
        
        proposal = accepted_cp.proposal
        student = candidature.student
        
        # Load the template
        template_path = '/app/templates/docs/protocol_template.docx'
        if not os.path.exists(template_path):
            return Response({"message": "Template de protocolo não encontrado"}, status=HTTP_500_INTERNAL_SERVER_ERROR)
        
        doc = Document(template_path)
        
        # Get calendar for dates
        calendar = student.calendar
        
        with transaction.atomic():
            # Create protocol record FIRST to get the protocol number
            academic_year = f"{calendar.calendar_year}/{calendar.calendar_year+1}" if calendar else 'N/A'
            protocol = Protocol.objects.create(
                candidature=candidature,
                academic_year=academic_year
            )
            
            # NOW prepare replacement data with actual protocol number
            replacements = {
                '{{protocol_number}}': protocol.protocol_number or '',
                '{{academic_year}}': academic_year,
                '{{company_name}}': proposal.company.company_name if proposal.company else 'ISEC',
                '{{student_name}}': student.student_name,
                '{{student_number}}': str(student.student_number),
                '{{course_name}}': proposal.course.course_name if proposal.course else 'N/A',
                '{{proposal_title}}': proposal.proposal_title or '',
                '{{proposal_description}}': proposal.proposal_description or '',
                '{{proposal_objectives}}': proposal.proposal_objectives or '',
                '{{proposal_location}}': proposal.location or 'A definir',
                '{{start_date}}': str(calendar.submission_start) if calendar else 'A definir',
                '{{end_date}}': str(calendar.placements) if calendar else 'A definir',
                '{{isec_advisor_name}}': proposal.isec_advisor.teacher_name if proposal.isec_advisor else 'A designar',
                '{{company_advisor_name}}': proposal.company_advisor.representative_name if proposal.company_advisor else 'A designar',
                '{{isec_signature_date}}': '___/___/______',
                '{{company_signature_date}}': '___/___/______',
                '{{student_signature_date}}': '___/___/______',
            }
            
            # Replace placeholders in all paragraphs
            for para in doc.paragraphs:
                for key, value in replacements.items():
                    if key in para.text:
                        for run in para.runs:
                            run.text = run.text.replace(key, value)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in para.text:
                                    for run in para.runs:
                                        run.text = run.text.replace(key, value)
            
            # Save document to bytes
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            # Save document file
            filename = f"protocol_{protocol.protocol_number}.docx"
            protocol.document.save(filename, ContentFile(doc_bytes.read()))
            
            # SIMPLIFICATION: Auto-sign ISEC when generating (admin is already doing it)
            protocol.isec_signed_at = timezone.now()
            protocol.isec_signed_by = account
            protocol.save()
            
            # Update candidature state - skip protocol_generated, go directly to awaiting company signature
            old_state = candidature.state
            candidature.state = 'awaiting_signatures'  # New simplified state
            candidature.save()
            
            # Log history
            CandidatureHistory.objects.create(
                candidature=candidature,
                previous_state=old_state,
                new_state='awaiting_signatures',
                changed_by=account,
                notes=f"Protocolo {protocol.protocol_number} gerado e assinado pelo ISEC"
            )
        
        return Response({
            "message": "Protocolo gerado e assinado pelo ISEC",
            "protocol_id": protocol.id_protocol,
            "protocol_number": protocol.protocol_number,
            "document_url": f"/api/protocol/{protocol.id_protocol}/download",
            "isec_signed": True
        }, status=HTTP_201_CREATED)
        
    except Candidature.DoesNotExist:
        return Response({"message": "Candidatura não encontrada"}, status=HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({"message": "Erro ao gerar protocolo", "details": str(e)}, status=HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(["GET"])
def downloadProtocol(request, pk):
    """
    Download protocol document.
    """
    auth_header = request.headers.get("Authorization")
    account, user_type, error = get_user_account(auth_header)
    
    if error:
        return Response({"message": "login"}, status=HTTP_401_UNAUTHORIZED)
    
    try:
        protocol = Protocol.objects.get(id_protocol=pk)
        
        # Check permissions - student can only download their own protocol
        if user_type == "student":
            student = Student.objects.get(user=account)
            if protocol.candidature.student != student:
                return Response({"message": "Não tem permissão para descarregar este protocolo"}, status=HTTP_403_FORBIDDEN)
        
        if not protocol.document:
            return Response({"message": "Documento de protocolo não encontrado"}, status=HTTP_404_NOT_FOUND)
        
        response = FileResponse(
            protocol.document.open('rb'),
            as_attachment=True,
            filename=f"{protocol.protocol_number}.docx"
        )
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response
        
    except Protocol.DoesNotExist:
        return Response({"message": "Protocolo não encontrado"}, status=HTTP_404_NOT_FOUND)
    except Student.DoesNotExist:
        return Response({"message": "Estudante não encontrado"}, status=HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({"message": "Erro ao descarregar protocolo", "details": str(e)}, status=HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(["GET"])
def getProtocol(request, pk):
    """
    Get protocol details.
    """
    auth_header = request.headers.get("Authorization")
    account, user_type, error = get_user_account(auth_header)
    
    if error:
        return Response({"message": "login"}, status=HTTP_401_UNAUTHORIZED)
    
    try:
        protocol = Protocol.objects.select_related(
            'candidature__student__user',
            'isec_signed_by',
            'company_signed_by'
        ).get(id_protocol=pk)
        
        # Check permissions - student can only view their own protocol
        if user_type == "student":
            student = Student.objects.get(user=account)
            if protocol.candidature.student != student:
                return Response({"message": "Não tem permissão para ver este protocolo"}, status=HTTP_403_FORBIDDEN)
        
        candidature = protocol.candidature
        student = candidature.student
        
        # Get accepted proposal
        accepted_cp = CandidatureProposal.objects.filter(
            candidature=candidature,
            state='accepted'
        ).select_related('proposal__company', 'proposal__course', 'proposal__isec_advisor').first()
        
        data = {
            "id": protocol.id_protocol,
            "protocol_number": protocol.protocol_number,
            "academic_year": protocol.academic_year,
            "generated_at": protocol.generated_at.isoformat() if protocol.generated_at else None,
            "candidature": {
                "id": candidature.id_candidature,
                "state": candidature.state,
                "student": {
                    "number": student.student_number,
                    "name": student.student_name,
                    "email": student.user.email
                }
            },
            "proposal": {
                "id": accepted_cp.proposal.id_proposal,
                "title": accepted_cp.proposal.proposal_title,
                "company": accepted_cp.proposal.company.company_name if accepted_cp.proposal.company else "ISEC",
                "course": accepted_cp.proposal.course.course_name if accepted_cp.proposal.course else "N/A"
            } if accepted_cp else None,
            "signatures": {
                "isec": {
                    "signed": protocol.isec_signed_at is not None,
                    "signed_at": protocol.isec_signed_at.isoformat() if protocol.isec_signed_at else None,
                    "signed_by": protocol.isec_signed_by.email if protocol.isec_signed_by else None
                },
                "company": {
                    "signed": protocol.company_signed_at is not None,
                    "signed_at": protocol.company_signed_at.isoformat() if protocol.company_signed_at else None,
                    "signed_by": protocol.company_signed_by.email if protocol.company_signed_by else None
                },
                "student": {
                    "signed": protocol.student_signed_at is not None,
                    "signed_at": protocol.student_signed_at.isoformat() if protocol.student_signed_at else None
                }
            },
            "has_document": bool(protocol.document)
        }
        
        return Response(data, status=HTTP_200_OK)
        
    except Protocol.DoesNotExist:
        return Response({"message": "Protocolo não encontrado"}, status=HTTP_404_NOT_FOUND)
    except Student.DoesNotExist:
        return Response({"message": "Estudante não encontrado"}, status=HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({"message": "Erro ao obter protocolo", "details": str(e)}, status=HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(["POST"])
def signProtocol(request, pk):
    """
    REQ-7: Sign protocol (ISEC, company, or student signature).
    Advances candidature through protocol workflow states.
    """
    auth_header = request.headers.get("Authorization")
    account, user_type, error = get_user_account(auth_header)
    
    if error:
        return Response({"message": "login"}, status=HTTP_401_UNAUTHORIZED)
    
    try:
        protocol = Protocol.objects.select_related('candidature').get(id_protocol=pk)
        candidature = protocol.candidature
        
        # SIMPLIFIED: Only company and student signatures needed (ISEC signs on generation)
        signature_type = request.data.get("signature_type")
        
        if not signature_type:
            # Auto-determine based on user type and what's missing
            if not protocol.company_signed_at and user_type in ['admin', 'representative']:
                signature_type = 'company'
            elif not protocol.student_signed_at and user_type in ['admin', 'student']:
                signature_type = 'student'
            else:
                return Response({"message": "Não há assinaturas pendentes para si"}, status=HTTP_400_BAD_REQUEST)
        
        with transaction.atomic():
            now = timezone.now()
            old_state = candidature.state
            
            if signature_type == 'company':
                if protocol.company_signed_at:
                    return Response({"message": "Empresa já assinou"}, status=HTTP_400_BAD_REQUEST)
                if user_type not in ['admin', 'representative']:
                    return Response({"message": "Apenas representantes de empresa podem assinar"}, status=HTTP_403_FORBIDDEN)
                
                protocol.company_signed_at = now
                protocol.company_signed_by = account
                
            elif signature_type == 'student':
                if protocol.student_signed_at:
                    return Response({"message": "Estudante já assinou"}, status=HTTP_400_BAD_REQUEST)
                if user_type not in ['admin', 'student']:
                    return Response({"message": "Apenas o estudante pode assinar"}, status=HTTP_403_FORBIDDEN)
                
                # Verify it's the right student
                if user_type == 'student':
                    student = Student.objects.get(user=account)
                    if candidature.student != student:
                        return Response({"message": "Apenas o estudante da candidatura pode assinar"}, status=HTTP_403_FORBIDDEN)
                
                protocol.student_signed_at = now
            
            else:
                return Response({"message": "Tipo de assinatura inválido"}, status=HTTP_400_BAD_REQUEST)
            
            protocol.save()
            
            # SIMPLIFICATION: Auto-finish when all signatures collected
            all_signed = (protocol.isec_signed_at and 
                         protocol.company_signed_at and 
                         protocol.student_signed_at)
            
            if all_signed:
                candidature.state = 'finished'
                candidature.save()
                CandidatureHistory.objects.create(
                    candidature=candidature,
                    previous_state=old_state,
                    new_state='finished',
                    changed_by=account,
                    notes=f"Protocolo completado - todas as assinaturas recolhidas"
                )
            else:
                CandidatureHistory.objects.create(
                    candidature=candidature,
                    previous_state=old_state,
                    new_state=candidature.state,
                    changed_by=account,
                    notes=f"Protocolo assinado ({signature_type})"
                )
        
        return Response({
            "message": "Protocolo assinado com sucesso",
            "signature_type": signature_type,
            "candidature_state": candidature.state,
            "all_signed": all_signed
        }, status=HTTP_200_OK)
        
    except Protocol.DoesNotExist:
        return Response({"message": "Protocolo não encontrado"}, status=HTTP_404_NOT_FOUND)
    except Student.DoesNotExist:
        return Response({"message": "Estudante não encontrado"}, status=HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({"message": "Erro ao assinar protocolo", "details": str(e)}, status=HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(["POST"])
def completeProtocol(request, pk):
    """
    Mark protocol/internship as finished.
    Only callable when all signatures are complete.
    """
    auth_header = request.headers.get("Authorization")
    account, user_type, error = get_user_account(auth_header)
    
    if error:
        return Response({"message": "login"}, status=HTTP_401_UNAUTHORIZED)
    
    if user_type not in ['admin', 'teacher']:
        return Response({"message": "Não tem permissão para finalizar protocolos"}, status=HTTP_403_FORBIDDEN)
    
    try:
        protocol = Protocol.objects.select_related('candidature').get(id_protocol=pk)
        candidature = protocol.candidature
        
        if candidature.state != 'student_signature':
            return Response(
                {"message": "Todas as assinaturas devem estar completas antes de finalizar"},
                status=HTTP_400_BAD_REQUEST
            )
        
        with transaction.atomic():
            old_state = candidature.state
            candidature.state = 'finished'
            candidature.save()
            
            # Log history
            CandidatureHistory.objects.create(
                candidature=candidature,
                previous_state=old_state,
                new_state='finished',
                changed_by=account,
                notes="Estágio/Projeto finalizado"
            )
        
        return Response({"message": "Protocolo finalizado com sucesso"}, status=HTTP_200_OK)
        
    except Protocol.DoesNotExist:
        return Response({"message": "Protocolo não encontrado"}, status=HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({"message": "Erro ao finalizar protocolo", "details": str(e)}, status=HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(["GET"])
def listProtocols(request):
    """
    List all protocols (with filters).
    """
    auth_header = request.headers.get("Authorization")
    account, user_type, error = get_user_account(auth_header)
    
    if error:
        return Response({"message": "login"}, status=HTTP_401_UNAUTHORIZED)
    
    try:
        # Base queryset
        protocols = Protocol.objects.select_related(
            'candidature__student__user',
            'candidature__student__calendar'
        ).all()
        
        # Filter by user type
        if user_type == "student":
            student = Student.objects.get(user=account)
            protocols = protocols.filter(candidature__student=student)
        elif user_type == "representative":
            # Filter by company
            from api.models import Representative
            rep = Representative.objects.get(user=account)
            protocols = protocols.filter(
                candidature__candidature_proposals__proposal__company=rep.company,
                candidature__candidature_proposals__state='accepted'
            ).distinct()
        
        # Apply query filters
        state = request.query_params.get('state')
        if state:
            protocols = protocols.filter(candidature__state=state)
        
        calendar_id = request.query_params.get('calendar')
        if calendar_id:
            protocols = protocols.filter(candidature__student__calendar__id_calendar=calendar_id)
        
        data = []
        for p in protocols:
            data.append({
                "id": p.id_protocol,
                "protocol_number": p.protocol_number,
                "academic_year": p.academic_year,
                "generated_at": p.generated_at.isoformat() if p.generated_at else None,
                "student": {
                    "number": p.candidature.student.student_number,
                    "name": p.candidature.student.student_name
                },
                "candidature_state": p.candidature.state,
                "signatures_complete": all([
                    p.isec_signed_at,
                    p.company_signed_at,
                    p.student_signed_at
                ])
            })
        
        return Response(data, status=HTTP_200_OK)
        
    except Student.DoesNotExist:
        return Response({"message": "Estudante não encontrado"}, status=HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({"message": "Erro ao listar protocolos", "details": str(e)}, status=HTTP_500_INTERNAL_SERVER_ERROR)

